#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
RAG向量数据库模块 - Milvus版本
支持Milvus向量数据库，用于存储和检索中医知识文档
无需下载HuggingFace模型，使用本地嵌入
"""

import os
import hashlib
import numpy as np
from typing import List, Dict, Optional, Tuple
from dataclasses import dataclass
import json

from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

# 尝试导入pymilvus
try:
    from pymilvus import (
        connections,
        utility,
        FieldSchema,
        CollectionSchema,
        DataType,
        Collection,
    )
    MILVUS_AVAILABLE = True
except ImportError:
    MILVUS_AVAILABLE = False
    print("[WARNING] pymilvus not installed, using mock implementation")


@dataclass
class SearchResult:
    """搜索结果数据类"""
    content: str
    metadata: Dict
    score: float
    source: str


class SimpleEmbedding:
    """
    简单本地嵌入模型
    使用词频统计生成向量，无需下载外部模型
    """
    
    def __init__(self, dim: int = 128):
        self.dim = dim
        # 中医领域常用词表
        self.vocab = self._build_vocab()
        
    def _build_vocab(self) -> Dict[str, int]:
        """构建中医领域词表"""
        # 常用中医词汇
        tcm_words = [
            # 经络穴位
            "经络", "穴位", "经脉", "气血", "阴阳", "五行", "脏腑", "十二经",
            "手太阴肺经", "手阳明大肠经", "足阳明胃经", "足太阴脾经",
            "手少阴心经", "手太阳小肠经", "足太阳膀胱经", "足少阴肾经",
            "手厥阴心包经", "手少阳三焦经", "足少阳胆经", "足厥阴肝经",
            "督脉", "任脉", "冲脉", "带脉",
            "足三里", "合谷", "太冲", "内关", "百会", "风池", "大椎",
            
            # 中药
            "中药", "草药", "人参", "黄芪", "当归", "白术", "茯苓", "甘草",
            "桂枝", "芍药", "生姜", "大枣", "附子", "干姜", "黄连", "黄芩",
            "解表", "清热", "温里", "补益", "理气", "活血", "化痰", "祛湿",
            
            # 方剂
            "方剂", "汤剂", "散剂", "丸剂", "桂枝汤", "麻黄汤", "小柴胡汤",
            "大承气汤", "四君子汤", "四物汤", "六味地黄丸", "金匮肾气丸",
            
            # 诊断
            "脉象", "舌象", "舌苔", "脉浮", "脉沉", "脉数", "脉迟",
            "表证", "里证", "寒证", "热证", "虚证", "实证",
            "风寒", "风热", "湿热", "痰湿", "气滞", "血瘀",
            
            # 症状
            "头痛", "发热", "恶寒", "汗出", "恶心", "呕吐", "腹痛",
            "胸闷", "心悸", "失眠", "咳嗽", "痰多", "便秘", "腹泻",
            
            # 经典
            "伤寒论", "金匮要略", "黄帝内经", "神农本草经", "温病条辨",
            "六经", "八纲", "辨证", "论治", "治则", "治法"
        ]
        
        vocab = {}
        for i, word in enumerate(tcm_words[:self.dim]):
            vocab[word] = i
        return vocab
    
    def embed_query(self, text: str) -> List[float]:
        """将文本转换为向量"""
        vector = np.zeros(self.dim)
        
        # 统计词频
        for word, idx in self.vocab.items():
            count = text.count(word)
            if count > 0:
                vector[idx] = count
        
        # 归一化
        norm = np.linalg.norm(vector)
        if norm > 0:
            vector = vector / norm
        
        return vector.tolist()
    
    def embed_documents(self, texts: List[str]) -> List[List[float]]:
        """批量转换文档为向量"""
        return [self.embed_query(text) for text in texts]


class TCMMilvusStore:
    """
    中医知识向量数据库 - Milvus实现
    用于存储和检索：针灸、伤寒论、金匮要略、神农本草经、医案等文档
    """
    
    def __init__(
        self,
        host: str = "localhost",
        port: str = "19530",
        collection_name: str = "tcm_knowledge",
        dim: int = 128,
        chunk_size: int = 512,
        chunk_overlap: int = 50,
        local_mode: bool = True  # 使用本地文件存储模式
    ):
        self.host = host
        self.port = port
        self.collection_name = collection_name
        self.dim = dim
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.local_mode = local_mode
        
        # 本地存储路径
        if local_mode:
            self.local_data_dir = "./data/milvus_local"
            os.makedirs(self.local_data_dir, exist_ok=True)
            self.local_data_file = os.path.join(self.local_data_dir, f"{collection_name}.json")
        
        # 初始化嵌入模型（本地，无需下载）
        self.embeddings = SimpleEmbedding(dim=dim)
        
        # 本地数据存储
        self.local_data: List[Dict] = []
        self._load_local_data()
        
        # 尝试连接Milvus（如果可用）
        self.milvus_collection = None
        if MILVUS_AVAILABLE and not local_mode:
            try:
                self._init_milvus()
            except Exception as e:
                print(f"[WARNING] Milvus connection failed: {e}, using local mode")
                self.local_mode = True
        
        # 文本分割器
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=["\n\n", "\n", "。", "；", "，", " ", ""],
            length_function=len
        )
    
    def _init_milvus(self):
        """初始化Milvus连接"""
        connections.connect(alias="default", host=self.host, port=self.port)
        
        # 检查集合是否存在
        if utility.has_collection(self.collection_name):
            self.milvus_collection = Collection(self.collection_name)
        else:
            # 创建集合
            self._create_collection()
    
    def _create_collection(self):
        """创建Milvus集合"""
        fields = [
            FieldSchema(name="id", dtype=DataType.VARCHAR, max_length=64, is_primary=True),
            FieldSchema(name="content", dtype=DataType.VARCHAR, max_length=65535),
            FieldSchema(name="source_type", dtype=DataType.VARCHAR, max_length=32),
            FieldSchema(name="embedding", dtype=DataType.FLOAT_VECTOR, dim=self.dim)
        ]
        
        schema = CollectionSchema(fields, description="TCM Knowledge Base")
        self.milvus_collection = Collection(self.collection_name, schema)
        
        # 创建索引
        index_params = {
            "metric_type": "COSINE",
            "index_type": "IVF_FLAT",
            "params": {"nlist": 128}
        }
        self.milvus_collection.create_index(field_name="embedding", index_params=index_params)
    
    def _load_local_data(self):
        """加载本地数据"""
        if self.local_mode and os.path.exists(self.local_data_file):
            try:
                with open(self.local_data_file, 'r', encoding='utf-8') as f:
                    self.local_data = json.load(f)
                print(f"[INFO] Loaded {len(self.local_data)} documents from local storage")
            except Exception as e:
                print(f"[WARNING] Failed to load local data: {e}")
                self.local_data = []
    
    def _save_local_data(self):
        """保存本地数据"""
        if self.local_mode:
            try:
                with open(self.local_data_file, 'w', encoding='utf-8') as f:
                    json.dump(self.local_data, f, ensure_ascii=False, indent=2)
            except Exception as e:
                print(f"[WARNING] Failed to save local data: {e}")
    
    def add_documents(
        self,
        documents: List[Document],
        source_type: str = "general"
    ) -> List[str]:
        """
        添加文档到向量数据库
        
        Args:
            documents: 文档列表
            source_type: 文档来源类型 (acupuncture/shanghan/jinkui/shennong/case)
        
        Returns:
            文档ID列表
        """
        # 分割文档
        chunks = self.text_splitter.split_documents(documents)
        
        ids = []
        for chunk in chunks:
            # 生成文档唯一ID
            content_hash = hashlib.md5(chunk.page_content.encode()).hexdigest()[:12]
            doc_id = f"{source_type}_{content_hash}"
            
            # 生成嵌入向量
            embedding = self.embeddings.embed_query(chunk.page_content)
            
            # 存储数据
            doc_data = {
                "id": doc_id,
                "content": chunk.page_content,
                "source_type": source_type,
                "embedding": embedding,
                "metadata": chunk.metadata
            }
            
            # 添加到本地存储
            if self.local_mode:
                # 检查是否已存在
                existing = [d for d in self.local_data if d["id"] == doc_id]
                if not existing:
                    self.local_data.append(doc_data)
            
            # 添加到Milvus（如果可用）
            if self.milvus_collection:
                self.milvus_collection.insert([
                    [doc_id],
                    [chunk.page_content],
                    [source_type],
                    [embedding]
                ])
            
            ids.append(doc_id)
        
        # 保存本地数据
        if self.local_mode:
            self._save_local_data()
        
        if self.milvus_collection:
            self.milvus_collection.flush()
        
        return ids
    
    def similarity_search(
        self,
        query: str,
        k: int = 5,
        source_type: Optional[str] = None
    ) -> List[SearchResult]:
        """
        相似度搜索
        
        Args:
            query: 查询文本
            k: 返回结果数量
            source_type: 筛选特定来源类型
        
        Returns:
            搜索结果列表
        """
        # 生成查询向量
        query_embedding = self.embeddings.embed_query(query)
        
        # 本地搜索
        if self.local_mode:
            return self._local_search(query_embedding, k, source_type)
        
        # Milvus搜索
        if self.milvus_collection:
            return self._milvus_search(query_embedding, k, source_type)
        
        return []
    
    def _local_search(
        self,
        query_embedding: List[float],
        k: int,
        source_type: Optional[str]
    ) -> List[SearchResult]:
        """本地数据搜索"""
        query_vec = np.array(query_embedding)
        
        # 筛选数据
        candidates = self.local_data
        if source_type:
            candidates = [d for d in candidates if d["source_type"] == source_type]
        
        # 计算相似度
        results = []
        for doc in candidates:
            doc_vec = np.array(doc["embedding"])
            # 余弦相似度
            similarity = np.dot(query_vec, doc_vec) / (
                np.linalg.norm(query_vec) * np.linalg.norm(doc_vec) + 1e-8
            )
            results.append((similarity, doc))
        
        # 排序并返回前k个
        results.sort(key=lambda x: x[0], reverse=True)
        
        return [
            SearchResult(
                content=doc["content"],
                metadata=doc.get("metadata", {}),
                score=float(score),
                source=doc["source_type"]
            )
            for score, doc in results[:k]
        ]
    
    def _milvus_search(
        self,
        query_embedding: List[float],
        k: int,
        source_type: Optional[str]
    ) -> List[SearchResult]:
        """Milvus搜索"""
        search_params = {"metric_type": "COSINE", "params": {"nprobe": 10}}
        
        # 构建过滤条件
        expr = None
        if source_type:
            expr = f'source_type == "{source_type}"'
        
        results = self.milvus_collection.search(
            data=[query_embedding],
            anns_field="embedding",
            param=search_params,
            limit=k,
            expr=expr,
            output_fields=["content", "source_type"]
        )
        
        search_results = []
        for hits in results:
            for hit in hits:
                search_results.append(SearchResult(
                    content=hit.entity.get("content"),
                    metadata={},
                    score=hit.score,
                    source=hit.entity.get("source_type")
                ))
        
        return search_results
    
    def search_by_source_type(
        self,
        query: str,
        source_type: str,
        k: int = 5
    ) -> List[SearchResult]:
        """
        按来源类型搜索
        
        Args:
            query: 查询文本
            source_type: 来源类型
            k: 返回结果数量
        
        Returns:
            搜索结果列表
        """
        return self.similarity_search(query, k, source_type)
    
    def multi_source_search(
        self,
        query: str,
        source_types: List[str],
        k_per_source: int = 3
    ) -> Dict[str, List[SearchResult]]:
        """
        多来源并行搜索
        
        Args:
            query: 查询文本
            source_types: 来源类型列表
            k_per_source: 每类返回数量
        
        Returns:
            按来源分类的搜索结果
        """
        results = {}
        for source_type in source_types:
            results[source_type] = self.similarity_search(
                query, k_per_source, source_type
            )
        return results
    
    def get_collection_stats(self) -> Dict:
        """获取集合统计信息"""
        if self.local_mode:
            return {
                "mode": "local",
                "total_documents": len(self.local_data),
                "collection_name": self.collection_name,
                "dimension": self.dim
            }
        
        if self.milvus_collection:
            stats = {
                "mode": "milvus",
                "collection_name": self.collection_name,
                "dimension": self.dim
            }
            try:
                stats["total_documents"] = self.milvus_collection.num_entities
            except:
                stats["total_documents"] = 0
            return stats
        
        return {"mode": "none", "total_documents": 0}
    
    def delete_collection(self):
        """删除集合"""
        if self.local_mode:
            self.local_data = []
            self._save_local_data()
        
        if self.milvus_collection and MILVUS_AVAILABLE:
            utility.drop_collection(self.collection_name)


class DocumentLoader:
    """文档加载器 - 支持多种格式"""
    
    SUPPORTED_EXTENSIONS = {
        '.txt', '.md', '.pdf', '.docx', '.json'
    }
    
    @classmethod
    def load_from_file(cls, file_path: str) -> Optional[Document]:
        """
        从文件加载文档
        
        Args:
            file_path: 文件路径
        
        Returns:
            Document对象或None
        """
        if not os.path.exists(file_path):
            return None
        
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext == '.txt' or ext == '.md':
            return cls._load_text(file_path)
        elif ext == '.json':
            return cls._load_json(file_path)
        elif ext == '.pdf':
            return cls._load_pdf(file_path)
        elif ext == '.docx':
            return cls._load_docx(file_path)
        else:
            print(f"[WARNING] Unsupported file type: {ext}")
            return None
    
    @classmethod
    def load_from_directory(
        cls,
        directory: str,
        source_type: str = "general"
    ) -> List[Document]:
        """
        从目录批量加载文档
        
        Args:
            directory: 目录路径
            source_type: 文档来源类型
        
        Returns:
            Document列表
        """
        documents = []
        
        if not os.path.exists(directory):
            return documents
        
        for filename in os.listdir(directory):
            file_path = os.path.join(directory, filename)
            if os.path.isfile(file_path):
                doc = cls.load_from_file(file_path)
                if doc:
                    doc.metadata["source"] = filename
                    doc.metadata["source_type"] = source_type
                    documents.append(doc)
        
        return documents
    
    @classmethod
    def _load_text(cls, file_path: str) -> Document:
        """加载文本文件"""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        return Document(page_content=content, metadata={"source": file_path})
    
    @classmethod
    def _load_json(cls, file_path: str) -> Document:
        """加载JSON文件"""
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        content = json.dumps(data, ensure_ascii=False, indent=2)
        return Document(page_content=content, metadata={"source": file_path})
    
    @classmethod
    def _load_pdf(cls, file_path: str) -> Optional[Document]:
        """加载PDF文件"""
        try:
            import PyPDF2
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                content = ""
                for page in reader.pages:
                    content += page.extract_text() + "\n"
            return Document(page_content=content, metadata={"source": file_path})
        except ImportError:
            print("[WARNING] PyPDF2 not installed, cannot load PDF")
            return None
    
    @classmethod
    def _load_docx(cls, file_path: str) -> Optional[Document]:
        """加载Word文件"""
        try:
            import docx
            doc = docx.Document(file_path)
            content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            return Document(page_content=content, metadata={"source": file_path})
        except ImportError:
            print("[WARNING] python-docx not installed, cannot load DOCX")
            return None


# 兼容性：保持与旧版相同的接口
TCMVectorStore = TCMMilvusStore
